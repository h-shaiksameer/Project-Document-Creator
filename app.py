import os
import time
import re
import requests
from dotenv import load_dotenv
from docx import Document
from googleapiclient.discovery import build
from docx.shared import Pt
import google.generativeai as genai
from google.api_core.exceptions import InternalServerError

# Load environment variables
load_dotenv()
api_key = os.getenv("API_KEY")
genai.configure(api_key=api_key)
gemini_model = genai.GenerativeModel("gemini-1.5-flash-8b")

# Function to get response from Gemini model with retry mechanism
def get_gemini_response(model, message):
    retry_attempts = 3  # Set the number of retry attempts
    for attempt in range(retry_attempts):
        try:
            chat = model.start_chat()
            response = chat.send_message(message)
            return response.text
        except InternalServerError as e:
            print(f"Attempt {attempt + 1} failed due to server error: {e}. Retrying...")
            time.sleep(2)  # Wait for 2 seconds before retrying
        except requests.exceptions.Timeout as e:
            print(f"Attempt {attempt + 1} failed due to timeout: {e}. Retrying...")
            time.sleep(2)  # Wait for 2 seconds before retrying
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            break  # Stop the loop on any unexpected error
    return "Failed to get a response after multiple attempts."

# Preprocess the content before adding to the document
def preprocess_content(content):
    # Remove lines with only special characters or introductory/closing statements
    content = re.sub(r'^[#*].*', '', content)  # Remove lines that start with '#' or '*'
    content = re.sub(r'^\s*(This (is|was).*)|(\(.*\))|(\*|#).*$', '', content)  # Remove unwanted lines
    content = re.sub(r'\n+', '\n', content).strip()  # Remove extra newlines
    return content

# Function to create the project document
def create_project_document(project_description):
    # Initialize the document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Define the table of contents
    table_of_contents = [
        "Abstract",
        "Table of Contents",
        "Introduction",
        "Literature Survey",
        "Analysis and Design",
        "Experimental Investigations",
        "Implementation",
        "Testing and Debugging/Results",
        "Conclusion / Bibliography",
        "References",
        "Appendices"
    ]

    # Add project description to the document
    doc.add_heading('Project Documentation', 0)
    doc.add_paragraph('Project Description:')
    doc.add_paragraph(project_description)

    # Function to generate content for each section
    def generate_section_content(section_title, prompt):
        doc.add_page_break()  # Add a page break for each new section
        # Add heading and make it bold
        heading = doc.add_heading(level=1)
        run = heading.add_run(section_title)
        run.bold = True
        run.font.size = Pt(17)
        # Generate content
        content = get_gemini_response(gemini_model, prompt)
        content = preprocess_content(content)
        # Add content to the document with normal text style
        doc.add_paragraph(content)

    # Generate content for the first page
    generate_section_content("Abstract", f"Generate an abstract for the project based on this description: {project_description}")
    generate_section_content("Table of Contents", f"Generate a table of contents for the project based on this description: {project_description}")
    generate_section_content("Introduction", f"Generate an introduction for the project based on this description: {project_description}")
    generate_section_content("Literature Survey", f"Generate a literature survey for the project based on this description: {project_description}")
    generate_section_content("Conclusion", f"Generate a conclusion for the project based on this description: {project_description}")
    generate_section_content("References", f"Generate references for the project based on this description: {project_description}")
    generate_section_content("Appendices", f"Generate appendices for the project based on this description: {project_description}")

    # Generate content for the remaining sections (these will be split across multiple pages if needed)
    generate_section_content("Analysis and Design", f"Generate analysis and design details for the project based on this description: {project_description}")
    generate_section_content("Experimental Investigations", f"Generate experimental investigations for the project based on this description: {project_description}")
    generate_section_content("Implementation", f"Generate implementation details for the project based on this description: {project_description}")
    generate_section_content("Testing and Debugging/Results", f"Generate testing and debugging/results for the project based on this description: {project_description}")

    # Save the document
    doc.save("project_documentation.docx")
    print("Document generated successfully!")

# Main function to get user input
def main():
    # Get project description from the user
    print("Enter the project description:")
    project_description = input()
    
    # Create the project document based on user input
    create_project_document(project_description)

if __name__ == "__main__":
    main()




# import os
# import time
# import re
# import requests
# from dotenv import load_dotenv
# from docx import Document
# from docx.oxml import parse_xml
# from docx.oxml.ns import nsdecls
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# import google.generativeai as genai
# from google.api_core.exceptions import InternalServerError

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("API_KEY")
# genai.configure(api_key=api_key)
# gemini_model = genai.GenerativeModel("gemini-1.5-flash-8b")

# # Function to get response from Gemini model with retry mechanism
# def get_gemini_response(model, message):
#     retry_attempts = 3  # Set the number of retry attempts
#     for attempt in range(retry_attempts):
#         try:
#             chat = model.start_chat()
#             response = chat.send_message(message)
#             return response.text
#         except InternalServerError as e:
#             print(f"Attempt {attempt + 1} failed due to server error: {e}. Retrying...")
#             time.sleep(2)  # Wait for 2 seconds before retrying
#         except requests.exceptions.Timeout as e:
#             print(f"Attempt {attempt + 1} failed due to timeout: {e}. Retrying...")
#             time.sleep(2)  # Wait for 2 seconds before retrying
#         except Exception as e:
#             print(f"An unexpected error occurred: {e}")
#             break  # Stop the loop on any unexpected error
#     return "Failed to get a response after multiple attempts."

# # Preprocess the content
# def preprocess_content(content):
#     # Remove lines with only special characters or introductory/closing statements
#     content = re.sub(r'^[#*].*', '', content)  # Remove lines that start with '#' or '*'
#     content = re.sub(r'^\s*(This (is|was).*)|(\(.*\))|(\*|#).*$', '', content)  # Remove unwanted lines
#     content = re.sub(r'\*\*', '', content)  # Remove bold markdown (**)
#     content = re.sub(r'\*', '', content)    # Remove other markdown characters (*)
#     content = re.sub(r'#', '', content)    # Remove markdown header symbols (#)
#     content = re.sub(r'\n+', '\n', content).strip()  # Remove extra newlines
#     return content

# # Check if content is a code block or link
# def handle_code_and_links(content, doc):
#     paragraphs = content.split("\n")
#     for paragraph in paragraphs:
#         if paragraph.startswith("http://") or paragraph.startswith("https://"):
#             # Add hyperlink
#             add_hyperlink(doc, paragraph, paragraph)
#         elif re.match(r'^\s*(\w+)\s*{.*', paragraph) or paragraph.strip().startswith("def "):
#             # Add code block
#             add_code_block(doc, paragraph)
#         else:
#             # Add normal text
#             doc.add_paragraph(paragraph)

# # Add a hyperlink to the document
# def add_hyperlink(doc, text, url):
#     p = doc.add_paragraph()
#     run = p.add_run(text)
#     run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
#     run.underline = True
#     r = run._element
#     hyperlink = parse_xml(
#         f'<w:hyperlink r:id="{url}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">{r.xml}</w:hyperlink>'
#     )
#     doc.element.body.append(hyperlink)

# # Add a code block to the document
# def add_code_block(doc, code):
#     p = doc.add_paragraph()
#     p.text = code
#     p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#     # Set the background to black and font to white
#     shading = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
#     p._p.get_or_add_pPr().append(shading)
#     for run in p.runs:
#         run.font.color.rgb = RGBColor(255, 255, 255)  # White color
#         run.font.name = "Courier New"  # Monospace font

# # Function to create the project document
# def create_project_document(project_description):
#     # Initialize the document
#     doc = Document()

#     # Set up styles
#     style = doc.styles['Normal']
#     style.font.name = 'Times New Roman'
#     style.font.size = Pt(14)

#     # Fixed Table of Contents
#     table_of_contents = [
#         "Abstract",
#         "Table of Contents",
#         "Introduction",
#         "Literature Survey",
#         "Analysis and Design",
#         "Experimental Investigations",
#         "Implementation",
#         "Testing and Debugging/Results",
#         "Conclusion / Bibliography",
#         "References",
#         "Appendices"
#     ]

#     # Add the fixed Table of Contents page
#     doc.add_heading('Project Documentation', 0)
#     doc.add_paragraph('Project Description:')
#     doc.add_paragraph(project_description)

#     doc.add_page_break()
#     doc.add_heading('Table of Contents', level=1)
#     for section in table_of_contents:
#         doc.add_paragraph(section)

#     # Generate content for each section
#     def generate_section_content(section_title, prompt):
#         doc.add_page_break()
#         heading = doc.add_heading(level=1)
#         run = heading.add_run(section_title)
#         run.bold = True
#         run.font.size = Pt(17)
#         content = get_gemini_response(gemini_model, prompt)
#         content = preprocess_content(content)
#         handle_code_and_links(content, doc)

#     # Generate each section
#     generate_section_content("Abstract", f"Generate an abstract for the project based on this description: {project_description}")
#     generate_section_content("Introduction", f"Generate an introduction for the project based on this description: {project_description}")
#     generate_section_content("Literature Survey", f"Generate a literature survey for the project based on this description: {project_description}")
#     generate_section_content("Analysis and Design", f"Generate analysis and design details for the project based on this description: {project_description}")
#     generate_section_content("Experimental Investigations", f"Generate experimental investigations for the project based on this description: {project_description}")
#     generate_section_content("Implementation", f"Generate implementation details for the project based on this description: {project_description}")
#     generate_section_content("Testing and Debugging/Results", f"Generate testing and debugging/results for the project based on this description: {project_description}")
#     generate_section_content("Conclusion / Bibliography", f"Generate conclusion and bibliography for the project based on this description: {project_description}")
#     generate_section_content("References", f"Generate references for the project based on this description: {project_description}")
#     generate_section_content("Appendices", f"Generate appendices for the project based on this description: {project_description}")

#     doc.save("project_documentation.docx")
#     print("Document generated successfully!")

# # Main function to get user input
# def main():
#     print("Enter the project description:")
#     project_description = input()
#     create_project_document(project_description)

# if __name__ == "__main__":
#     main()
