import os
import time
import re
import requests
from dotenv import load_dotenv
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
from google.api_core.exceptions import InternalServerError
from absl import logging
logging.set_verbosity(logging.INFO)

os.environ["GRPC_VERBOSITY"] = "ERROR"

# Load environment variables
load_dotenv()
api_key = os.getenv("API_KEY")
genai.configure(api_key=api_key)
gemini_model = genai.GenerativeModel("gemini-1.5-flash-8b")

# Function to get response from Gemini model with retry mechanism
def get_gemini_response(model, message):
    retry_attempts = 3
    for attempt in range(retry_attempts):
        try:
            chat = model.start_chat()
            response = chat.send_message(message)
            return response.text
        except InternalServerError as e:
            print(f"Attempt {attempt + 1} failed due to server error: {e}. Retrying...")
            time.sleep(2)
        except requests.exceptions.Timeout as e:
            print(f"Attempt {attempt + 1} failed due to timeout: {e}. Retrying...")
            time.sleep(2)
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            break
    return "Failed to get a response after multiple attempts."

# Preprocess the content
def preprocess_content(content):
    # Remove lines with only special characters or introductory/closing statements
    content = re.sub(r'^[#*].*', '', content)  # Remove lines that start with '#' or '*'
    content = re.sub(r'^\s*(This (is|was).*)|(\(.*\))|(\*|#).*$', '', content)  # Remove unwanted lines
    content = re.sub(r'\*\*', '', content)  # Remove bold markdown (**)
    content = re.sub(r'\*', '', content)    # Remove other markdown characters (*)
    content = re.sub(r'#', '', content)    # Remove markdown header symbols (#)
    content = re.sub(r'\n+', '\n', content).strip()  # Remove extra newlines
    return content

# Check if content is a code block or link
def handle_code_and_links(content, doc):
    paragraphs = content.split("\n")
    for paragraph in paragraphs:
        if paragraph.startswith("http://") or paragraph.startswith("https://"):
            # Add hyperlink
            add_hyperlink(doc, paragraph, paragraph)
        elif re.match(r'^\s*(\w+)\s*{.*', paragraph) or paragraph.strip().startswith("def "):
            # Add code block
            add_code_block(doc, paragraph)
        else:
            # Add normal text
            doc.add_paragraph(paragraph)

# Add a hyperlink to the document
def add_hyperlink(doc, text, url):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
    run.underline = True
    p.add_hyperlink(url)

# Add a code block to the document
def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.text = code
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shading = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
    p._p.get_or_add_pPr().append(shading)
    for run in p.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)  # White color
        run.font.name = "Courier New"  # Monospace font

# Generate a sanitized filename from the project description
def generate_filename(project_description):
    # Extract the first 3-4 words from the description
    words = project_description.split()[:4]
    base_name = "_".join(words)
    # Remove invalid characters for filenames
    sanitized_name = re.sub(r'[\\/*?:"<>|]', '', base_name)
    return f"{sanitized_name}_documentation.docx"

def create_table_of_contents(doc):
    # Add a heading for the table of contents
    toc_heading = doc.add_heading(level=0)  # Use level 0 for the main heading
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(26)  # Set font size for the TOC heading
    
    # Add a table with bold borders for the TOC entries
    table = doc.add_table(rows=0, cols=2)  # Create a table with two columns (Section Name and Page Number)
    table.style = 'Table Grid'  # Bold borders for the table
    
    sections = [
        "Abstract",
        "Table of Contents",
        "Introduction",
        "Literature Survey",
        "Analysis and Design",
        "Experimental Investigations",
        "Implementation",
        "Testing and Debugging/Results",
        "Conclusion / Bibliography",
        "References",
        "Appendices"
    ]
    
    for section in sections:
        row = table.add_row()
        row.cells[0].text = section
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(22)  # Set font size for section names
        row.cells[1].text = "..."  # Placeholder for page numbers; could be updated dynamically later
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(22)


# Function to create the project document
def create_project_document(project_description):
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Fixed Table of Contents
    table_of_contents = [
        "Abstract",
        "Table of Contents",
        "Introduction",
        "Literature Survey",
        "Analysis and Design",
        "Experimental Investigations",
        "Implementation",
        "Testing and Debugging/Results",
        "Conclusion / Bibliography",
        "References",
        "Appendices"
    ]

    doc.add_heading('Project Documentation', 0)
    doc.add_paragraph('Project Description:')
    doc.add_paragraph(project_description)

    doc.add_page_break()
    create_table_of_contents(doc)

    def generate_section_content(section_title, prompt):
        doc.add_page_break()  # Start each section on a new page
        heading = doc.add_heading(level=1)
        run = heading.add_run(section_title)
        run.bold = True
        run.font.size = Pt(17)
        
        # Generate content
        content = get_gemini_response(gemini_model, prompt)
        content = preprocess_content(content)
        
        # Process references or appendices differently if needed
        if section_title == "References":
            for reference in content.split("\n"):
                if reference.strip():
                    doc.add_paragraph(reference.strip(), style="List Bullet")
        elif section_title == "Appendices":
            # Treat appendices as subsections
            for appendix in content.split("\n\n"):  # Assuming sections are separated by double newlines
                if appendix.strip():
                    doc.add_heading(appendix.split(":")[0], level=2)  # Subsection title
                    doc.add_paragraph(":".join(appendix.split(":")[1:]).strip())
        else:
            # Add content for other sections
            doc.add_paragraph(content)


    # Generate content for each section
    generate_section_content("Abstract", f"Generate an abstract for the project based on this description: {project_description}")
    generate_section_content(
        "Introduction",
        f"Write an engaging introduction for the project titled '{project_description}'. The introduction should clearly explain the purpose, scope, and goals of the project. It should be accessible to a wide audience and avoid excessive jargon."
    )

    generate_section_content(
        "Literature Survey",
        f"Conduct a detailed literature survey on relevant technologies and methodologies related to the project titled '{project_description}'. The survey should cover:\n1. Core technologies used in the field of the project.\n2. Research trends, innovations, and advancements in the domain.\n3. Case studies or relevant examples.\n4. Ethical considerations and challenges.\n5. Future opportunities in this field."
    )

    generate_section_content("Analysis and Design", f"Provide a detailed analysis and design for the project titled '{project_description}'. This should include a breakdown of major components, system architecture, and any design patterns or frameworks used.")
    generate_section_content("Experimental Investigations", f"Describe any experimental investigations or studies conducted for the project titled '{project_description}'. This could include testing methods, data collection, and analysis.")
    generate_section_content(
        "Implementation",
        f"Provide detailed implementation steps for building the project titled '{project_description}'. Include a breakdown of the core components, architecture, tools, and technologies used in the development process."
    )

    generate_section_content(
        "Testing and Debugging/Results",
        f"Create a comprehensive testing and debugging strategy for the project titled '{project_description}'. This should cover various testing approaches (functional, performance, security) and debugging techniques used during development."
    )

    generate_section_content("Conclusion / Bibliography", f"Summarize the key findings, achievements, and limitations of the project titled '{project_description}'. Include recommendations for future work, followed by a bibliography listing relevant references.")
    
    generate_section_content("References", f"Generate a list of references related to the project titled '{project_description}'. Provide references to relevant books, articles, research papers, websites, and tools used.")
    
    generate_section_content("Appendices", f"Create appendices for the project titled '{project_description}', including technical specifications, detailed data, and additional resources that support the project.")

    # Generate filename and save the document
    filename = generate_filename(project_description)
    doc.save(filename)
    print(f"Document saved as {filename}")

# Main function to get user input
def main():
    print("Enter the project description:")
    project_description = input()
    create_project_document(project_description)

if __name__ == "__main__":
    main()
